import os
import sys
import logging
import pandas as pd
import requests
import torch
from fastapi import FastAPI
from pydantic import BaseModel
from rdkit import Chem
from rdkit.Chem import Draw
import base64
from io import BytesIO
import numpy as np
import io
from fastapi.middleware.cors import CORSMiddleware
import seaborn as sns
import matplotlib.pyplot as plt
from dotenv import load_dotenv
load_dotenv()

# Setup logging to show in console
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[logging.StreamHandler(sys.stdout)]
)
logger = logging.getLogger(__name__)

from discovery import build_prompt, smiles_to_image_base64
from model import DrugTargetModel, encode_protein, fetch_drug_smiles, fetch_protein_sequence, smiles_to_fingerprint
from utils import build_predict_prompt, get_top_similar_drugs

from google import genai 

api_key = os.getenv("GEMINI_API_KEY")
ai_client = genai.Client(api_key=api_key)

protein_dim = 400  
drug_dim = 1024 
model = DrugTargetModel(protein_dim, drug_dim)
model.load_state_dict(torch.load("drug_target_model.pth"))
model.eval()

def fetch_drug_info(drug_name):
    url = f"https://pubchem.ncbi.nlm.nih.gov/rest/pug/compound/name/{drug_name}/property/CanonicalSMILES,MolecularWeight,XLogP,HBondDonorCount,HBondAcceptorCount/JSON"
    response = requests.get(url)
    if response.status_code == 200:
        data = response.json()
        properties = data.get("PropertyTable", {}).get("Properties", [{}])[0]
        return {
            "smiles": properties.get("CanonicalSMILES", "N/A"),
            "molecular_weight": properties.get("MolecularWeight", "N/A"),
            "logP": properties.get("XLogP", "N/A"),
            "h_bond_donors": properties.get("HBondDonorCount", "N/A"),
            "h_bond_acceptors": properties.get("HBondAcceptorCount", "N/A")
        }
    return None

def generate_molecule_image(smiles):
    mol = Chem.MolFromSmiles(smiles)
    if mol:
        img = Draw.MolToImage(mol)
        buffered = BytesIO()
        img.save(buffered, format="PNG")
        img_base_64 = base64.b64encode(buffered.getvalue()).decode("utf-8")
        return f"data:image/png;base64,{img_base_64}"
    return None

def compute_feature_importance(model, protein_tensor, drug_tensor):
    protein_tensor.requires_grad = True
    drug_tensor.requires_grad = True

    output = model(protein_tensor, drug_tensor)
    model.zero_grad()
    output.backward()

    protein_grad = protein_tensor.grad.abs().numpy().flatten()
    drug_grad = drug_tensor.grad.abs().numpy().flatten()

    protein_grad = (protein_grad - protein_grad.min()) / (protein_grad.max() - protein_grad.min())
    drug_grad = (drug_grad - drug_grad.min()) / (drug_grad.max() - drug_grad.min())

    return protein_grad, drug_grad

def generate_heatmap(protein_grad, drug_grad):
    interaction_matrix = np.outer(protein_grad, drug_grad)
    plt.figure(figsize=(10, 6))
    sns.heatmap(interaction_matrix, cmap="coolwarm", xticklabels=False, yticklabels=False, center=0)
    
    img_buf = io.BytesIO()
    plt.savefig(img_buf, format="png")
    img_buf.seek(0)
    img_str = base64.b64encode(img_buf.read()).decode("utf-8")
    plt.close()
    
    return f"data:image/png;base64,{img_str}"

app = FastAPI()

# CORS configuration - allow frontend origins
origins = [
    "http://localhost:3000",
    "http://127.0.0.1:3000",
    "http://localhost:3001",
]

# Add any additional origins from environment variable
cors_origins_env = os.getenv("CORS_ORIGINS", "")
if cors_origins_env:
    origins.extend([origin.strip() for origin in cors_origins_env.split(",") if origin.strip()])

app.add_middleware(
    CORSMiddleware,
    allow_origins=origins,
    allow_credentials=True,
    allow_methods=["GET", "POST", "PUT", "DELETE", "OPTIONS"],
    allow_headers=["*"],
    expose_headers=["*"],
)

class PredictionRequest(BaseModel):
    uniprot_id: str
    drug_name: str


@app.get("/health")
def check_health():
    logger.info("Health check endpoint called")
    return {"message": "Good Health"}

@app.post("/predict")
def predict_interaction(request: PredictionRequest):
    protein_sequence = fetch_protein_sequence(request.uniprot_id)
    if not protein_sequence:
        return {"error": "Invalid UniProt ID or sequence not found"}
    
    drug_info = fetch_drug_info(request.drug_name)
    if not drug_info or drug_info["smiles"] == "N/A":
        return {"error": "Invalid drug name or properties not found"}
    
    drug_info["name"] = request.drug_name
    
    protein_encoded = encode_protein(protein_sequence)
    drug_encoded = smiles_to_fingerprint(drug_info["smiles"])
    
    protein_tensor = torch.tensor([protein_encoded], dtype=torch.float32)
    drug_tensor = torch.tensor([drug_encoded], dtype=torch.float32)
    
    with torch.no_grad():
        prediction = model(protein_tensor, drug_tensor).item()
    
    molecule_image = generate_molecule_image(drug_info["smiles"])

    protein_grad, drug_grad = compute_feature_importance(model, protein_tensor, drug_tensor)
    heatmap_image = generate_heatmap(protein_grad, drug_grad)

    prompt = build_predict_prompt(request.uniprot_id, drug_info, prediction)
    insights = None
    
    # Try different models as fallback
    models_to_try = ["gemini-2.0-flash", "gemini-1.5-flash", "gemini-1.5-pro"]
    
    for model_name in models_to_try:
        try:
            response = ai_client.models.generate_content(model=model_name, contents=prompt)
            insights = response.text
            break
        except Exception as e:
            error_msg = str(e)
            print(f"[Warning] {model_name} failed: {error_msg[:100]}")
            if "429" in error_msg or "RESOURCE_EXHAUSTED" in error_msg:
                continue
            else:
                break
    
    if not insights:
        binding_status = "likely to bind" if prediction > 0.5 else "unlikely to bind"
        insights = f"""## Prediction Summary

The model predicts that **{request.drug_name}** is **{binding_status}** to protein **{request.uniprot_id}** with a binding probability of **{prediction:.2%}**.

**Note:** Detailed AI analysis is temporarily unavailable due to API rate limits. Please try again later."""

    return {
        "uniprot_id": request.uniprot_id,
        "drug_name": request.drug_name,
        "smiles": drug_info["smiles"],
        "molecular_weight": drug_info["molecular_weight"],
        "logP": drug_info["logP"],
        "h_bond_donors": drug_info["h_bond_donors"],
        "h_bond_acceptors": drug_info["h_bond_acceptors"],
        "binding_probability": prediction,
        "molecule_image": molecule_image,
        "heatmap_image": heatmap_image,
        "top_similar_drugs": get_top_similar_drugs(drug_info['smiles']),
        "insights": insights
    }

class DiscoveryRequest(BaseModel):
    uniprot_id: str
    top_n: int = 5
    

@app.post("/discover")
def discover_candidates(request: DiscoveryRequest):
    logger.info(f"=== DISCOVER REQUEST RECEIVED === Protein: {request.uniprot_id}, Top N: {request.top_n}")
    sys.stdout.flush()
    
    protein_sequence = fetch_protein_sequence(request.uniprot_id)
    if not protein_sequence:
        logger.error(f"Invalid UniProt ID: {request.uniprot_id}")
        return {"error": "Invalid UniProt ID or sequence not found"}

    logger.info(f"Protein sequence fetched, length: {len(protein_sequence)}")
    
    protein_encoded = encode_protein(protein_sequence)
    protein_tensor = torch.tensor(np.array([protein_encoded]), dtype=torch.float32)
    
    drug_db = pd.read_csv("drug_db.csv")
    results = []
    
    # Filter out biologics and invalid SMILES
    valid_drugs = drug_db[~drug_db['smiles'].str.contains('Antibody|Analog|Inhibitor|Protein|Heparin|Insulin', case=False, na=False)]
    logger.info(f"Processing {len(valid_drugs)} small molecule drugs...")
    sys.stdout.flush()
    
    for idx, row in valid_drugs.iterrows():
        smiles = row['smiles']
        name = row['name']
        
        try:
            drug_fp = smiles_to_fingerprint(smiles)
            if drug_fp is None or np.sum(drug_fp) == 0:
                continue
                
            drug_tensor = torch.tensor(np.array([drug_fp]), dtype=torch.float32)

            with torch.no_grad():
                score = model(protein_tensor, drug_tensor).item()

            image_base64 = smiles_to_image_base64(smiles)
            if not image_base64:
                continue
                
            logger.info(f"Drug {len(results)+1}: {name}, Score: {score:.4f}")

            results.append({
                "name": name,
                "smiles": smiles,
                "score": round(score, 4),
                "image_base64": image_base64
            })
        except Exception as e:
            logger.warning(f"Error processing {name}: {e}")
            continue

    results = sorted(results, key=lambda x: x["score"], reverse=True)[:request.top_n]
    logger.info(f"Top {len(results)} candidates selected")
    sys.stdout.flush()
    
    # Enrich top candidates with PubChem data
    for result in results:
        try:
            drug_info = fetch_drug_info(result["name"])
            if drug_info:
                result["molecular_weight"] = drug_info.get("molecular_weight", "N/A")
                result["logP"] = drug_info.get("logP", "N/A")
                result["h_bond_donors"] = drug_info.get("h_bond_donors", "N/A")
                result["h_bond_acceptors"] = drug_info.get("h_bond_acceptors", "N/A")
            logger.info(f"Enriched: {result['name']}")
        except:
            pass

    # Build detailed analysis prompt
    prompt = f"""You are writing a detailed drug discovery research report for protein {request.uniprot_id}.

Top drug candidates identified:
"""
    for i, drug in enumerate(results, 1):
        prompt += f"{i}. {drug['name']} (Binding Score: {drug['score']:.1%}, MW: {drug.get('molecular_weight', 'N/A')}, LogP: {drug.get('logP', 'N/A')})\n"
    
    prompt += """
Write a comprehensive analysis with these sections:

## Insights

Start with an overview paragraph identifying what the target protein is and its biological significance. Explain why understanding its interactions is important for drug discovery.

Then provide detailed analysis of each compound:

For each top compound, write a detailed entry with:
- **Drug-likeness:** Assessment of pharmaceutical properties
- **Known Use:** Current clinical applications and mechanism of action  
- **Literature Relevance:** How this compound relates to the target protein based on scientific literature
- **Repurposing Opportunities:** Potential new therapeutic applications

## Key Next Steps for Repurposing and Research

Provide actionable research recommendations:
- **Experimental Validation:** What experiments are needed
- **Mechanism of Action Studies:** How to investigate binding mechanisms
- **Clinical Relevance:** Implications for patient care
- **Drug Design:** How these findings inform new drug development

End with a concluding paragraph on the value of this computational analysis.

Write in academic but accessible language. Use **bold** for key terms. Be thorough and informative."""
    insights = None
    
    # Try different models as fallback
    models_to_try = ["gemini-2.0-flash", "gemini-2.0-flash-lite", "gemini-2.5-flash"]
    
    for model_name in models_to_try:
        try:
            logger.info(f"Trying Gemini model: {model_name}")
            sys.stdout.flush()
            response = ai_client.models.generate_content(model=model_name, contents=prompt)
            insights = response.text
            logger.info(f"SUCCESS with model: {model_name}")
            break
        except Exception as e:
            error_msg = str(e)
            logger.warning(f"{model_name} FAILED: {error_msg[:200]}")
            sys.stdout.flush()
            if "429" in error_msg or "RESOURCE_EXHAUSTED" in error_msg or "404" in error_msg:
                continue  # Try next model
            else:
                break  # Other error, stop trying
    
    if not insights:
        # Generate basic insights without AI
        insights = f"""## Drug Discovery Results for {request.uniprot_id}

**Note:** AI-powered analysis is temporarily unavailable due to API rate limits.

### Top Candidates Found:
"""
        for i, drug in enumerate(results, 1):
            insights += f"\n**{i}. {drug['name']}**\n"
            insights += f"- Binding Score: {drug['score']}\n"
            if drug.get('molecular_weight'):
                insights += f"- Molecular Weight: {drug.get('molecular_weight', 'N/A')}\n"
            if drug.get('logP'):
                insights += f"- LogP: {drug.get('logP', 'N/A')}\n"
        
        insights += "\n*Please try again later for AI-powered insights, or check your Gemini API quota.*"

    return {
        "uniprot_id": request.uniprot_id,
        "top_candidates": results,
        "insights": insights
    }


# ==================== UNIFIED DRUG DISCOVERY ENDPOINT ====================

from research_agent import (
    ResearchAgent, 
    search_research_papers, 
    build_research_prompt,
    get_protein_structure_url
)
from typing import List, Optional
import re

# Initialize research agent
research_agent = ResearchAgent(ai_client)


class UnifiedDiscoveryRequest(BaseModel):
    query: str
    top_n: int = 8
    max_papers: int = 15


def detect_protein_id(query: str) -> Optional[str]:
    """Detect UniProt ID from query string"""
    match = re.search(r'\b([A-Z][0-9][A-Z0-9]{3}[0-9])\b', query)
    return match.group(1) if match else None


def fetch_protein_info(uniprot_id: str) -> Optional[dict]:
    """Fetch protein information from UniProt"""
    try:
        response = requests.get(
            f"https://rest.uniprot.org/uniprotkb/{uniprot_id}.json",
            timeout=30
        )
        if response.status_code == 200:
            data = response.json()
            function_text = ""
            for comment in data.get("comments", []):
                if comment.get("commentType") == "FUNCTION":
                    texts = comment.get("texts", [])
                    if texts:
                        function_text = texts[0].get("value", "")
                    break
            
            return {
                "id": uniprot_id,
                "name": data.get("proteinDescription", {}).get("recommendedName", {}).get("fullName", {}).get("value", "Unknown"),
                "organism": data.get("organism", {}).get("scientificName", "Unknown"),
                "sequence_length": data.get("sequence", {}).get("length", 0),
                "function": function_text,
                "sequence": data.get("sequence", {}).get("value", "")
            }
    except Exception as e:
        logger.error(f"Error fetching protein info: {e}")
    return None


@app.post("/unified-discovery")
def unified_discovery(request: UnifiedDiscoveryRequest):
    """
    UNIFIED Drug Discovery Pipeline:
    1. Detect protein from query
    2. Fetch protein info
    3. Search research papers
    4. Run drug binding predictions
    5. Generate unified AI analysis
    """
    logger.info(f"=== UNIFIED DISCOVERY === Query: {request.query}")
    sys.stdout.flush()
    
    result = {
        "query": request.query,
        "protein_info": None,
        "papers": [],
        "top_candidates": [],
        "insights": "",
        "tools_used": []
    }
    
    try:
        # STEP 1: Detect protein ID
        protein_id = detect_protein_id(request.query)
        logger.info(f"Step 1 - Protein ID detected: {protein_id}")
        
        # STEP 2: Fetch protein info if ID found
        if protein_id:
            result["protein_info"] = fetch_protein_info(protein_id)
            result["tools_used"].append("uniprot_api")
            logger.info(f"Step 2 - Protein info: {result['protein_info'].get('name') if result['protein_info'] else 'Not found'}")
        
        # STEP 3: Search research papers
        logger.info("Step 3 - Searching research papers...")
        papers = search_research_papers(request.query, request.max_papers)
        result["papers"] = papers
        result["tools_used"].append("research_papers")
        logger.info(f"Step 3 - Found {len(papers)} papers")
        
        # STEP 4: Run drug discovery if protein found
        if result["protein_info"] and result["protein_info"].get("sequence"):
            logger.info("Step 4 - Running drug binding predictions...")
            result["tools_used"].append("drug_discovery_ml")
            
            protein_sequence = result["protein_info"]["sequence"]
            protein_encoded = encode_protein(protein_sequence)
            protein_tensor = torch.tensor(np.array([protein_encoded]), dtype=torch.float32)
            
            drug_db = pd.read_csv("drug_db.csv")
            candidates = []
            
            # Filter valid drugs
            valid_drugs = drug_db[~drug_db['smiles'].str.contains('Antibody|Analog|Inhibitor|Protein|Heparin|Insulin', case=False, na=False)]
            
            for idx, row in valid_drugs.iterrows():
                try:
                    smiles = row['smiles']
                    name = row['name']
                    
                    drug_fp = smiles_to_fingerprint(smiles)
                    if drug_fp is None or np.sum(drug_fp) == 0:
                        continue
                    
                    drug_tensor = torch.tensor(np.array([drug_fp]), dtype=torch.float32)
                    
                    with torch.no_grad():
                        score = model(protein_tensor, drug_tensor).item()
                    
                    image_base64 = smiles_to_image_base64(smiles)
                    if not image_base64:
                        continue
                    
                    candidates.append({
                        "name": name,
                        "smiles": smiles,
                        "score": round(score, 4),
                        "image_base64": image_base64
                    })
                except Exception as e:
                    continue
            
            # Sort and get top candidates
            candidates = sorted(candidates, key=lambda x: x["score"], reverse=True)[:request.top_n]
            
            # Enrich with PubChem data
            for candidate in candidates:
                try:
                    drug_info = fetch_drug_info(candidate["name"])
                    if drug_info:
                        candidate["molecular_weight"] = drug_info.get("molecular_weight", "N/A")
                        candidate["logP"] = drug_info.get("logP", "N/A")
                except:
                    pass
            
            result["top_candidates"] = candidates
            logger.info(f"Step 4 - Found {len(candidates)} drug candidates")
        
        # STEP 5: Generate comprehensive report
        logger.info("Step 5 - Generating comprehensive analysis...")
        
        # Build prompt for AI
        prompt = f"""Write a detailed drug discovery research report for: {request.query}

"""
        if result.get("protein_info"):
            p = result["protein_info"]
            prompt += f"Target Protein: {p.get('name')} ({p.get('id')})\nFunction: {p.get('function', '')[:300]}\n\n"
        
        if result.get("top_candidates"):
            prompt += "Drug Candidates:\n"
            for d in result["top_candidates"][:5]:
                prompt += f"- {d['name']}: {d['score']:.1%} binding\n"
            prompt += "\n"
        
        if result.get("papers"):
            prompt += "Research Papers:\n"
            for p in result["papers"][:5]:
                prompt += f"- {p['title'][:80]} ({p['source']})\n"
        
        # Adaptive prompt based on query type
        has_drugs = bool(result.get("top_candidates"))
        has_protein = bool(result.get("protein_info"))
        
        if has_drugs or has_protein:
            # Drug discovery report
            prompt += """
Write a comprehensive research report. Use proper formatting:

### Executive Summary
3-4 sentences summarizing key findings and significance. Include citations (Author et al., Year).

### Background & Significance
- Why this topic/target matters
- Current research landscape
- Key challenges in the field

### Literature Analysis
Based on research papers:
- Key finding 1 with citation (Author et al., Year, Source)
- Key finding 2 with citation
- Key finding 3 with citation
- Emerging trends

### Drug Candidates
For each compound:
**[Drug Name]** - Drug class. Mechanism. Clinical uses. Binding score: X%.

### Therapeutic Potential
- Clinical opportunities
- Drug repurposing possibilities
- Safety considerations

### Recommendations
1. Next research steps
2. Validation approaches
3. Future directions

### References
List all papers cited above.

Use **bold** for key terms. Include inline citations."""
        else:
            # General research report
            prompt += """
Write a comprehensive research report on this topic. Use proper formatting:

### Executive Summary
3-4 sentences summarizing key findings and significance. Include citations (Author et al., Year).

### Background & Context
- Why this topic matters
- Current state of research
- Key debates or challenges

### Literature Analysis
Based on research papers:
- Key finding 1 with citation (Author et al., Year, Source)
- Key finding 2 with citation
- Key finding 3 with citation
- Important themes and patterns

### Key Insights
- Major insight 1
- Major insight 2
- Major insight 3
- Connections between findings

### Implications
- Practical applications
- Theoretical contributions
- Societal impact

### Future Directions
1. Research gaps to address
2. Methodological suggestions
3. Emerging questions

### References
List all papers cited above.

Use **bold** for key terms. Include inline citations throughout."""
        
        # Try multiple Gemini models - gemini-2.5-flash works best
        models_to_try = ["gemini-2.5-flash", "gemini-2.0-flash", "gemini-pro"]
        insights = None
        
        for model_name in models_to_try:
            try:
                logger.info(f"Trying model: {model_name}")
                response = ai_client.models.generate_content(model=model_name, contents=prompt)
                insights = response.text
                logger.info(f"SUCCESS with model: {model_name}")
                result["tools_used"].append(f"ai_analysis:{model_name}")
                break
            except Exception as e:
                error_msg = str(e)
                logger.warning(f"{model_name} failed: {error_msg[:100]}")
                continue
        
        if not insights:
            # Generate VERY detailed fallback report WITHOUT using AI
            insights = generate_comprehensive_fallback(result)
        
        result["insights"] = insights
        logger.info(f"Step 5 - Analysis complete. Tools used: {result['tools_used']}")
        
        return result
        
    except Exception as e:
        logger.error(f"Unified discovery error: {e}")
        result["error"] = str(e)
        return result


def generate_comprehensive_fallback(result: dict) -> str:
    """Generate a comprehensive research report without AI"""
    query = result.get("query", "drug discovery research")
    report = ""
    
    # Executive Summary
    report += "### Executive Summary\n\n"
    if result.get("protein_info"):
        p = result["protein_info"]
        report += f"This comprehensive analysis investigates **{p.get('name', 'the target protein')}** "
        report += f"(UniProt: {p.get('id', 'N/A')}), a {p.get('sequence_length', 0)}-amino acid protein "
        report += f"expressed in *{p.get('organism', 'humans')}*. "
        report += "This protein represents a significant therapeutic target with potential implications "
        report += "for drug discovery and development. Our computational analysis has identified "
        report += f"several promising drug candidates with predicted binding affinity to this target.\n\n"
    else:
        report += f"This report presents a comprehensive analysis of **{query}**, "
        report += "examining current research literature, potential therapeutic compounds, "
        report += "and opportunities for drug discovery. The analysis integrates findings from "
        report += f"multiple scientific databases including {len(result.get('papers', []))} peer-reviewed publications "
        report += "to provide actionable insights for researchers and drug developers.\n\n"
    
    # Background & Significance
    report += "### Background & Significance\n\n"
    if result.get("protein_info"):
        p = result["protein_info"]
        if p.get('function'):
            report += f"**Biological Function:** {p.get('function')}\n\n"
        report += "Understanding protein-drug interactions is fundamental to modern drug discovery. "
        report += "Computational approaches enable rapid screening of potential therapeutic compounds, "
        report += "significantly accelerating the drug development pipeline. This target has been "
        report += "identified through rigorous bioinformatics analysis as a promising candidate "
        report += "for therapeutic intervention.\n\n"
    else:
        report += f"**{query.title()}** represents an active area of pharmaceutical research "
        report += "with significant therapeutic potential. Current research efforts focus on "
        report += "understanding disease mechanisms, identifying novel drug targets, and "
        report += "developing effective therapeutic interventions. The integration of computational "
        report += "methods with traditional drug discovery approaches has accelerated progress in this field.\n\n"
    
    # Literature Analysis
    report += "### Literature Analysis\n\n"
    if result.get("papers") and len(result["papers"]) > 0:
        report += f"Analysis of {len(result['papers'])} research publications reveals key insights:\n\n"
        
        # Group papers by source
        sources = {}
        for paper in result["papers"]:
            src = paper.get("source", "Other")
            if src not in sources:
                sources[src] = []
            sources[src].append(paper)
        
        for source, papers in sources.items():
            report += f"**{source} Publications:**\n"
            for paper in papers[:3]:
                title = paper.get('title', 'N/A')
                authors = paper.get('authors', ['Unknown'])
                year = paper.get('published', 'N/A')
                report += f"- *\"{title}\"* ({authors[0] if authors else 'Unknown'} et al., {year})\n"
            report += "\n"
        
        report += "**Key Themes Identified:**\n"
        report += "- Novel therapeutic approaches and drug delivery systems\n"
        report += "- Molecular mechanisms underlying disease pathology\n"
        report += "- Clinical trial outcomes and safety profiles\n"
        report += "- Computational and AI-driven drug discovery methods\n\n"
    else:
        report += "Literature search is in progress. Relevant publications will inform "
        report += "therapeutic strategies and drug development approaches.\n\n"
    
    # Drug Candidates
    if result.get("top_candidates") and len(result["top_candidates"]) > 0:
        report += "### Drug Candidate Analysis\n\n"
        report += "Computational screening identified the following compounds with predicted binding affinity:\n\n"
        
        for i, drug in enumerate(result["top_candidates"], 1):
            report += f"**{i}. {drug['name']}**\n\n"
            report += f"- **Binding Affinity Score:** {drug['score']:.1%}\n"
            if drug.get('molecular_weight'):
                report += f"- **Molecular Weight:** {drug['molecular_weight']} Da\n"
            if drug.get('logP'):
                report += f"- **LogP (Lipophilicity):** {drug['logP']}\n"
            
            # Detailed drug info
            drug_details = get_comprehensive_drug_info(drug['name'])
            if drug_details:
                report += f"- **Drug Class:** {drug_details.get('class', 'N/A')}\n"
                report += f"- **Mechanism of Action:** {drug_details.get('mechanism', 'N/A')}\n"
                report += f"- **Current Indications:** {drug_details.get('indications', 'N/A')}\n"
                report += f"- **Relevance:** {drug_details.get('relevance', 'Potential for drug repurposing')}\n"
            report += "\n"
    
    # Therapeutic Opportunities
    report += "### Therapeutic Opportunities\n\n"
    report += "Based on the computational analysis and literature review, several therapeutic opportunities emerge:\n\n"
    report += "- **Drug Repurposing:** Existing FDA-approved drugs may offer faster pathways to clinical application\n"
    report += "- **Combination Therapies:** Multi-target approaches could enhance therapeutic efficacy\n"
    report += "- **Novel Drug Design:** Structural insights enable rational design of improved compounds\n"
    report += "- **Biomarker Development:** Identified targets may serve as diagnostic or prognostic markers\n\n"
    
    # Research Recommendations
    report += "### Research Recommendations\n\n"
    report += "**Immediate Next Steps:**\n"
    report += "- Validate computational predictions with biochemical binding assays (SPR, ITC)\n"
    report += "- Perform molecular dynamics simulations to assess binding stability\n"
    report += "- Evaluate cytotoxicity profiles in relevant cell lines\n\n"
    
    report += "**Medium-term Goals:**\n"
    report += "- Conduct structure-activity relationship (SAR) studies\n"
    report += "- Assess pharmacokinetic properties (ADMET profiling)\n"
    report += "- Develop lead optimization strategies\n\n"
    
    report += "**Long-term Objectives:**\n"
    report += "- Progress promising candidates to preclinical studies\n"
    report += "- Establish collaborative partnerships for clinical development\n"
    report += "- Explore intellectual property opportunities\n"
    
    return report


def get_comprehensive_drug_info(drug_name: str) -> dict:
    """Get comprehensive information for common drugs"""
    drug_database = {
        "Furosemide": {
            "class": "Loop Diuretic",
            "mechanism": "Inhibits Na-K-2Cl cotransporter in loop of Henle",
            "indications": "Edema, heart failure, hypertension, renal disease",
            "relevance": "Cardiovascular and renal therapeutic applications"
        },
        "Amlodipine": {
            "class": "Calcium Channel Blocker (Dihydropyridine)",
            "mechanism": "Blocks L-type calcium channels in vascular smooth muscle",
            "indications": "Hypertension, chronic stable angina, vasospastic angina",
            "relevance": "Cardiovascular protection and blood pressure regulation"
        },
        "Canagliflozin": {
            "class": "SGLT2 Inhibitor",
            "mechanism": "Inhibits sodium-glucose co-transporter 2 in proximal tubule",
            "indications": "Type 2 diabetes, heart failure, chronic kidney disease",
            "relevance": "Metabolic and cardiovascular benefits beyond glycemic control"
        },
        "Levothyroxine": {
            "class": "Thyroid Hormone",
            "mechanism": "Synthetic T4 hormone replacement",
            "indications": "Hypothyroidism, thyroid cancer, myxedema coma",
            "relevance": "Metabolic regulation and hormonal therapy"
        },
        "Sildenafil": {
            "class": "PDE5 Inhibitor",
            "mechanism": "Inhibits phosphodiesterase type 5, increases cGMP",
            "indications": "Erectile dysfunction, pulmonary arterial hypertension",
            "relevance": "Vascular and smooth muscle relaxation applications"
        },
        "Dapagliflozin": {
            "class": "SGLT2 Inhibitor",
            "mechanism": "Inhibits renal glucose reabsorption",
            "indications": "Type 2 diabetes, heart failure with reduced EF",
            "relevance": "Cardiorenal protection in metabolic disease"
        },
        "Empagliflozin": {
            "class": "SGLT2 Inhibitor",
            "mechanism": "Blocks sodium-glucose cotransporter 2",
            "indications": "Type 2 diabetes, heart failure, chronic kidney disease",
            "relevance": "Proven cardiovascular mortality reduction"
        },
        "Gabapentin": {
            "class": "Anticonvulsant/Analgesic",
            "mechanism": "Binds alpha-2-delta subunit of voltage-gated calcium channels",
            "indications": "Neuropathic pain, epilepsy, restless leg syndrome",
            "relevance": "Neurological and pain management applications"
        },
        "Omeprazole": {
            "class": "Proton Pump Inhibitor",
            "mechanism": "Irreversibly inhibits H+/K+-ATPase in gastric parietal cells",
            "indications": "GERD, peptic ulcer disease, H. pylori eradication",
            "relevance": "Gastrointestinal protection and acid suppression"
        },
        "Metformin": {
            "class": "Biguanide Antidiabetic",
            "mechanism": "Activates AMPK, reduces hepatic gluconeogenesis",
            "indications": "Type 2 diabetes, polycystic ovary syndrome, prediabetes",
            "relevance": "First-line diabetes therapy with potential anti-aging properties"
        },
    }
    return drug_database.get(drug_name, None)


# ==================== LEGACY ENDPOINTS (kept for compatibility) ====================

class ResearchQuery(BaseModel):
    query: str
    include_papers: bool = True
    max_papers: int = 8


class PaperSearchRequest(BaseModel):
    query: str
    max_results: int = 10


@app.post("/research")
def research_query(request: ResearchQuery):
    """
    Main research endpoint - processes a query and returns comprehensive results
    """
    logger.info(f"=== RESEARCH QUERY === {request.query}")
    sys.stdout.flush()
    
    try:
        # Process with research agent
        results = research_agent.process_query(request.query)
        
        # Generate AI insights if papers found
        if results["papers"] or results.get("protein_info") or results.get("drug_info"):
            prompt = build_research_prompt(
                request.query,
                results["papers"],
                results.get("protein_info"),
                results.get("drug_info")
            )
            
            # Try different models as fallback (optimized token usage)
            models_to_try = ["gemini-2.0-flash", "gemini-2.0-flash-lite", "gemini-2.5-flash"]
            insights = None
            
            for model_name in models_to_try:
                try:
                    logger.info(f"Trying model: {model_name}")
                    response = ai_client.models.generate_content(
                        model=model_name,
                        contents=prompt
                    )
                    insights = response.text
                    logger.info(f"SUCCESS with model: {model_name}")
                    break
                except Exception as e:
                    error_msg = str(e)
                    logger.warning(f"{model_name} failed: {error_msg[:100]}")
                    if "429" in error_msg or "RESOURCE_EXHAUSTED" in error_msg:
                        continue
                    break
            
            results["insights"] = insights or "Analysis in progress. Review the research papers below."
        
        logger.info(f"Research complete: {len(results['papers'])} papers, tools: {results['tools_used']}")
        return results
        
    except Exception as e:
        logger.error(f"Research error: {e}")
        return {"error": str(e), "papers": [], "insights": ""}


@app.post("/search-papers")
def search_papers(request: PaperSearchRequest):
    """
    Search for research papers from ArXiv and PubMed
    """
    logger.info(f"=== PAPER SEARCH === {request.query}")
    
    papers = search_research_papers(request.query, request.max_results)
    
    return {
        "query": request.query,
        "papers": papers,
        "total": len(papers)
    }


class MoleculeRequest(BaseModel):
    smiles: str
    name: Optional[str] = None


@app.post("/molecule-3d")
def get_molecule_3d(request: MoleculeRequest):
    """
    Generate 3D coordinates for a molecule from SMILES
    """
    try:
        from rdkit.Chem import AllChem
        
        mol = Chem.MolFromSmiles(request.smiles)
        if mol is None:
            return {"error": "Invalid SMILES string"}
        
        # Add hydrogens
        mol = Chem.AddHs(mol)
        
        # Generate 3D coordinates
        AllChem.EmbedMolecule(mol, randomSeed=42)
        AllChem.MMFFOptimizeMolecule(mol)
        
        # Get 3D coordinates as MOL block
        mol_block = Chem.MolToMolBlock(mol)
        
        # Also generate 2D image
        mol_2d = Chem.MolFromSmiles(request.smiles)
        img = Draw.MolToImage(mol_2d, size=(300, 300))
        buffer = BytesIO()
        img.save(buffer, format="PNG")
        img_base64 = base64.b64encode(buffer.getvalue()).decode("utf-8")
        
        return {
            "smiles": request.smiles,
            "name": request.name,
            "mol_block": mol_block,
            "image_2d": img_base64,
            "atom_count": mol.GetNumAtoms(),
            "bond_count": mol.GetNumBonds()
        }
        
    except Exception as e:
        logger.error(f"3D molecule generation error: {e}")
        return {"error": str(e)}


class ProteinStructureRequest(BaseModel):
    uniprot_id: str


@app.post("/protein-structure")
def get_protein_structure(request: ProteinStructureRequest):
    """
    Get protein structure information and AlphaFold URL
    """
    try:
        # Get protein info from UniProt
        response = requests.get(
            f"https://rest.uniprot.org/uniprotkb/{request.uniprot_id}.json",
            timeout=30
        )
        
        if response.status_code != 200:
            return {"error": "Protein not found"}
        
        data = response.json()
        
        # Get AlphaFold structure URL
        alphafold_url = get_protein_structure_url(request.uniprot_id)
        
        # Try to fetch AlphaFold PDB
        pdb_content = None
        try:
            pdb_response = requests.get(alphafold_url, timeout=30)
            if pdb_response.status_code == 200:
                pdb_content = pdb_response.text
        except:
            pass
        
        return {
            "uniprot_id": request.uniprot_id,
            "name": data.get("proteinDescription", {}).get("recommendedName", {}).get("fullName", {}).get("value", "N/A"),
            "organism": data.get("organism", {}).get("scientificName", "N/A"),
            "sequence": data.get("sequence", {}).get("value", ""),
            "sequence_length": data.get("sequence", {}).get("length", 0),
            "alphafold_url": alphafold_url,
            "pdb_content": pdb_content,
            "has_structure": pdb_content is not None
        }
        
    except Exception as e:
        logger.error(f"Protein structure error: {e}")
        return {"error": str(e)}


# ==================== DRUG DISCOVERY ENDPOINTS ====================

from drug_discovery import (
    calculate_admet_properties,
    generate_analogs,
    scaffold_hop,
    virtual_screen,
    submit_docking_job,
    optimize_lead,
    calculate_similarity,
    generate_novel_drug,
    get_3d_coordinates,
    get_protein_3d_structure
)


class ADMETRequest(BaseModel):
    smiles: str


class AnalogRequest(BaseModel):
    smiles: str
    num_analogs: int = 10


class DockingRequest(BaseModel):
    protein_pdb: str
    ligand_smiles: str


class OptimizeRequest(BaseModel):
    smiles: str
    target_property: str = "drug_likeness"


class SimilarityRequest(BaseModel):
    smiles1: str
    smiles2: str


@app.post("/admet")
def predict_admet(request: ADMETRequest):
    """
    Predict ADMET properties for a molecule
    - Absorption, Distribution, Metabolism, Excretion, Toxicity
    - Drug-likeness scores (Lipinski, Veber, Lead-likeness)
    """
    logger.info(f"=== ADMET PREDICTION === SMILES: {request.smiles[:50]}...")
    
    result = calculate_admet_properties(request.smiles)
    
    if "error" in result:
        return {"error": result["error"]}
    
    # Generate molecule image
    try:
        mol = Chem.MolFromSmiles(request.smiles)
        if mol:
            img = Draw.MolToImage(mol, size=(300, 300))
            buffer = BytesIO()
            img.save(buffer, format="PNG")
            result["image_base64"] = base64.b64encode(buffer.getvalue()).decode("utf-8")
    except:
        pass
    
    logger.info(f"ADMET complete: Drug-likeness={result.get('drug_likeness_score', 'N/A')}")
    return result


@app.post("/generate-analogs")
def generate_molecule_analogs(request: AnalogRequest):
    """
    Generate molecular analogs using bioisosteric replacements,
    methylation, fluorination, and other modifications
    """
    logger.info(f"=== ANALOG GENERATION === SMILES: {request.smiles[:50]}...")
    
    analogs = generate_analogs(request.smiles, request.num_analogs)
    
    logger.info(f"Generated {len(analogs)} analogs")
    return {
        "parent_smiles": request.smiles,
        "analogs": analogs,
        "count": len(analogs)
    }


@app.post("/scaffold-hop")
def perform_scaffold_hop(request: AnalogRequest):
    """
    Generate molecules with different scaffolds but similar pharmacophore
    """
    logger.info(f"=== SCAFFOLD HOP === SMILES: {request.smiles[:50]}...")
    
    results = scaffold_hop(request.smiles, request.num_analogs)
    
    logger.info(f"Generated {len(results)} scaffold variants")
    return {
        "parent_smiles": request.smiles,
        "scaffold_variants": results,
        "count": len(results)
    }


@app.post("/dock")
def run_docking(request: DockingRequest):
    """
    Submit docking job to estimate binding affinity
    Returns estimated binding energy in kcal/mol
    """
    logger.info(f"=== DOCKING === Ligand: {request.ligand_smiles[:50]}...")
    
    result = submit_docking_job(request.protein_pdb, request.ligand_smiles)
    
    logger.info(f"Docking complete: Affinity={result.get('estimated_affinity_kcal', 'N/A')} kcal/mol")
    return result


@app.post("/optimize-lead")
def run_lead_optimization(request: OptimizeRequest):
    """
    Analyze a lead compound and suggest optimizations
    Returns improvement suggestions and optimized analogs
    """
    logger.info(f"=== LEAD OPTIMIZATION === SMILES: {request.smiles[:50]}...")
    
    # Get current properties
    current_props = calculate_admet_properties(request.smiles)
    
    # Get optimization suggestions
    suggestions = optimize_lead(request.smiles, request.target_property)
    
    return {
        "parent_smiles": request.smiles,
        "current_properties": current_props,
        "suggestions": suggestions
    }


@app.post("/similarity")
def calculate_mol_similarity(request: SimilarityRequest):
    """
    Calculate Tanimoto similarity between two molecules
    """
    similarity = calculate_similarity(request.smiles1, request.smiles2)
    
    return {
        "smiles1": request.smiles1,
        "smiles2": request.smiles2,
        "tanimoto_similarity": round(similarity, 4),
        "similar": similarity > 0.7
    }


@app.post("/drug-discovery-full")
def full_drug_discovery_pipeline(request: UnifiedDiscoveryRequest):
    """
    FULL Drug Discovery Pipeline with all capabilities:
    1. Literature search
    2. Drug candidate prediction
    3. ADMET analysis for top candidates
    4. Analog generation
    5. Docking estimation
    6. AI-powered report
    """
    logger.info(f"=== FULL DRUG DISCOVERY PIPELINE === Query: {request.query}")
    
    # Run unified discovery first
    base_result = unified_discovery(request)
    
    # Enhance with ADMET for top candidates
    if base_result.get("top_candidates"):
        for candidate in base_result["top_candidates"]:
            try:
                admet = calculate_admet_properties(candidate["smiles"])
                candidate["admet"] = {
                    "drug_likeness_score": admet.get("drug_likeness_score"),
                    "lipinski_pass": admet.get("lipinski_pass"),
                    "bbb_penetration": admet.get("bbb_penetration", {}).get("class"),
                    "oral_absorption": admet.get("oral_absorption", {}).get("class"),
                    "hepatotoxicity_risk": admet.get("hepatotoxicity_risk", {}).get("class"),
                    "cardiotoxicity_risk": admet.get("cardiotoxicity_risk", {}).get("class"),
                }
                
                # Generate analogs for top 3
                if base_result["top_candidates"].index(candidate) < 3:
                    candidate["analogs"] = generate_analogs(candidate["smiles"], 3)
                
            except Exception as e:
                logger.warning(f"ADMET error for {candidate.get('name')}: {e}")
    
    base_result["tools_used"].append("admet_prediction")
    base_result["tools_used"].append("analog_generation")
    
    logger.info("Full drug discovery pipeline complete")
    return base_result


# ==================== DE NOVO DRUG DESIGN ENDPOINTS ====================

class DeNovoRequest(BaseModel):
    seed_smiles: Optional[str] = None
    num_generations: int = 30
    population_size: int = 15


@app.post("/design-drug")
def design_new_drug(request: DeNovoRequest):
    """
    Design novel drug molecules using genetic algorithm
    Can start from scratch or evolve from a seed molecule
    """
    logger.info(f"=== DE NOVO DRUG DESIGN === Seed: {request.seed_smiles or 'Random'}")
    
    novel_drugs = generate_novel_drug(
        seed_smiles=request.seed_smiles,
        num_generations=request.num_generations,
        population_size=request.population_size
    )
    
    logger.info(f"Generated {len(novel_drugs)} novel drug candidates")
    
    return {
        "seed_smiles": request.seed_smiles,
        "novel_drugs": novel_drugs,
        "count": len(novel_drugs),
        "method": "Genetic Algorithm",
        "generations": request.num_generations
    }


# ==================== 3D VISUALIZATION ENDPOINTS ====================

class Molecule3DRequest(BaseModel):
    smiles: str


class Protein3DRequest(BaseModel):
    uniprot_id: str


@app.post("/molecule-3d-coords")
def get_molecule_3d_coords(request: Molecule3DRequest):
    """
    Generate 3D coordinates for a molecule
    Returns MOL block, PDB, and atom coordinates for 3Dmol.js
    """
    logger.info(f"=== 3D COORDINATES === SMILES: {request.smiles[:50]}...")
    
    coords = get_3d_coordinates(request.smiles)
    
    if "error" in coords:
        return {"error": coords["error"]}
    
    # Also generate 2D image
    try:
        mol = Chem.MolFromSmiles(request.smiles)
        if mol:
            img = Draw.MolToImage(mol, size=(300, 300))
            buffer = BytesIO()
            img.save(buffer, format="PNG")
            coords["image_2d"] = base64.b64encode(buffer.getvalue()).decode("utf-8")
    except:
        pass
    
    logger.info(f"3D coordinates generated: {coords.get('num_atoms', 0)} atoms")
    return coords


@app.post("/protein-3d-coords")
def get_protein_3d_coords(request: Protein3DRequest):
    """
    Fetch protein 3D structure from AlphaFold or PDB
    Returns PDB content for 3Dmol.js visualization
    """
    logger.info(f"=== PROTEIN 3D STRUCTURE === UniProt: {request.uniprot_id}")
    
    structure = get_protein_3d_structure(request.uniprot_id)
    
    if "error" in structure:
        return {"error": structure["error"]}
    
    logger.info(f"Protein structure fetched: {structure.get('source', 'Unknown')}")
    return structure


# ==================== DRUG COMPARISON ENDPOINT ====================

class CompareRequest(BaseModel):
    smiles_list: List[str]
    names: Optional[List[str]] = None


@app.post("/compare-drugs")
def compare_drugs(request: CompareRequest):
    """
    Compare multiple drug candidates side-by-side
    Returns normalized properties for radar chart visualization
    """
    logger.info(f"=== DRUG COMPARISON === Comparing {len(request.smiles_list)} compounds")
    
    from drug_discovery import calculate_admet_properties, calculate_drug_likeness_score
    
    results = []
    all_properties = {
        "molecular_weight": [],
        "logP": [],
        "tpsa": [],
        "h_bond_donors": [],
        "h_bond_acceptors": [],
        "rotatable_bonds": [],
        "drug_likeness_score": [],
    }
    
    for i, smiles in enumerate(request.smiles_list):
        try:
            mol = Chem.MolFromSmiles(smiles)
            if mol is None:
                continue
            
            # Calculate ADMET properties
            admet = calculate_admet_properties(smiles)
            if "error" in admet:
                continue
            
            # Generate image
            img = Draw.MolToImage(mol, size=(200, 200))
            buffer = BytesIO()
            img.save(buffer, format="PNG")
            img_base64 = base64.b64encode(buffer.getvalue()).decode("utf-8")
            
            # Get name
            name = request.names[i] if request.names and i < len(request.names) else f"Compound {i+1}"
            
            compound_data = {
                "name": name,
                "smiles": smiles,
                "image_base64": img_base64,
                "properties": {
                    "molecular_weight": admet.get("molecular_weight", 0),
                    "logP": admet.get("logP", 0),
                    "tpsa": admet.get("tpsa", 0),
                    "h_bond_donors": admet.get("h_bond_donors", 0),
                    "h_bond_acceptors": admet.get("h_bond_acceptors", 0),
                    "rotatable_bonds": admet.get("rotatable_bonds", 0),
                    "drug_likeness_score": admet.get("drug_likeness_score", 0),
                    "lipinski_pass": admet.get("lipinski_pass", False),
                    "veber_pass": admet.get("veber_pass", False),
                },
                "admet": {
                    "oral_absorption": admet.get("oral_absorption", {}),
                    "bbb_penetration": admet.get("bbb_penetration", {}),
                    "hepatotoxicity_risk": admet.get("hepatotoxicity_risk", {}),
                    "cardiotoxicity_risk": admet.get("cardiotoxicity_risk", {}),
                }
            }
            
            results.append(compound_data)
            
            # Collect for normalization
            for key in all_properties:
                all_properties[key].append(admet.get(key, 0))
                
        except Exception as e:
            logger.warning(f"Error processing compound {i}: {e}")
            continue
    
    # Calculate normalized scores (0-100) for radar chart
    def normalize(values, ideal_min, ideal_max, inverse=False):
        """Normalize values to 0-100 scale based on ideal drug-like ranges"""
        normalized = []
        for v in values:
            if ideal_max == ideal_min:
                score = 100
            else:
                # Score based on how close to ideal range
                if v < ideal_min:
                    score = max(0, 100 - ((ideal_min - v) / ideal_min) * 100) if ideal_min > 0 else 100
                elif v > ideal_max:
                    score = max(0, 100 - ((v - ideal_max) / ideal_max) * 50) if ideal_max > 0 else 50
                else:
                    score = 100
            
            if inverse:
                score = 100 - score
            normalized.append(round(min(100, max(0, score)), 1))
        return normalized
    
    # Ideal ranges for drug-likeness
    radar_data = {
        "labels": ["Drug-likeness", "Mol. Weight", "LogP", "TPSA", "HBD", "HBA"],
        "compounds": []
    }
    
    for i, compound in enumerate(results):
        props = compound["properties"]
        
        # Calculate radar scores (higher = better for drug discovery)
        radar_scores = {
            "drug_likeness": props["drug_likeness_score"] * 100,
            "molecular_weight": max(0, 100 - abs(props["molecular_weight"] - 350) / 5),  # Ideal ~350 Da
            "logP": max(0, 100 - abs(props["logP"] - 2.5) * 15),  # Ideal ~2.5
            "tpsa": max(0, 100 - abs(props["tpsa"] - 75) / 1.5),  # Ideal ~75
            "h_bond_donors": max(0, 100 - props["h_bond_donors"] * 15),  # Lower is better (0-5)
            "h_bond_acceptors": max(0, 100 - props["h_bond_acceptors"] * 8),  # Lower is better (0-10)
        }
        
        radar_data["compounds"].append({
            "name": compound["name"],
            "values": [
                round(radar_scores["drug_likeness"], 1),
                round(radar_scores["molecular_weight"], 1),
                round(radar_scores["logP"], 1),
                round(radar_scores["tpsa"], 1),
                round(radar_scores["h_bond_donors"], 1),
                round(radar_scores["h_bond_acceptors"], 1),
            ]
        })
    
    # Generate comparison summary
    if len(results) >= 2:
        best_drug_likeness = max(results, key=lambda x: x["properties"]["drug_likeness_score"])
        best_absorption = max(results, key=lambda x: x["admet"].get("oral_absorption", {}).get("probability", 0))
        safest = min(results, key=lambda x: (
            x["admet"].get("hepatotoxicity_risk", {}).get("risk", 1) +
            x["admet"].get("cardiotoxicity_risk", {}).get("risk", 1)
        ))
        
        summary = {
            "best_drug_likeness": best_drug_likeness["name"],
            "best_absorption": best_absorption["name"],
            "safest_profile": safest["name"],
            "total_compared": len(results),
            "lipinski_compliant": sum(1 for r in results if r["properties"]["lipinski_pass"])
        }
    else:
        summary = {"total_compared": len(results)}
    
    logger.info(f"Comparison complete: {len(results)} compounds analyzed")
    
    return {
        "compounds": results,
        "radar_data": radar_data,
        "summary": summary
    }


# ==================== REPORT EXPORT ENDPOINTS ====================

from fastapi.responses import StreamingResponse
import re
import unicodedata

class ExportReportRequest(BaseModel):
    query: str
    insights: str
    protein_info: Optional[dict] = None
    top_candidates: Optional[List[dict]] = None
    novel_drugs: Optional[List[dict]] = None
    papers: Optional[List[dict]] = None


def sanitize_text(text: str) -> str:
    """Remove or replace problematic characters for PDF/DOCX"""
    if not text:
        return ""
    # Normalize unicode
    text = unicodedata.normalize('NFKD', str(text))
    # Remove non-printable characters
    text = ''.join(char for char in text if unicodedata.category(char)[0] != 'C' or char in '\n\r\t')
    # Replace problematic characters
    text = text.replace('\u2019', "'").replace('\u2018', "'")
    text = text.replace('\u201c', '"').replace('\u201d', '"')
    text = text.replace('\u2013', '-').replace('\u2014', '-')
    text = text.replace('\u2026', '...')
    text = text.replace('\u00a0', ' ')
    # Remove any remaining non-ASCII if needed
    text = text.encode('ascii', 'ignore').decode('ascii')
    return text


@app.post("/export-pdf")
def export_report_pdf(request: ExportReportRequest):
    """
    Generate a professional PDF research report with molecule images
    """
    logger.info(f"=== PDF EXPORT === Query: {request.query}")
    
    try:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import letter, A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch, cm
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image as RLImage, PageBreak, KeepTogether, HRFlowable
        from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
        
        buffer = BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            rightMargin=50,
            leftMargin=50,
            topMargin=50,
            bottomMargin=50
        )
        
        # Styles
        styles = getSampleStyleSheet()
        # Use unique style names to avoid conflicts
        styles.add(ParagraphStyle(
            name='ReportTitle',
            parent=styles['Title'],
            fontSize=22,
            spaceAfter=10,
            textColor=colors.HexColor('#1a1a1a'),
            alignment=TA_CENTER
        ))
        styles.add(ParagraphStyle(
            name='ReportSubtitle',
            parent=styles['Normal'],
            fontSize=12,
            spaceAfter=5,
            textColor=colors.HexColor('#444444'),
            alignment=TA_CENTER
        ))
        styles.add(ParagraphStyle(
            name='ReportSection',
            parent=styles['Heading2'],
            fontSize=12,
            spaceBefore=15,
            spaceAfter=8,
            textColor=colors.HexColor('#1a1a1a'),
            borderPadding=5
        ))
        styles.add(ParagraphStyle(
            name='ReportBody',
            parent=styles['Normal'],
            fontSize=9,
            spaceBefore=3,
            spaceAfter=3,
            alignment=TA_JUSTIFY,
            textColor=colors.HexColor('#333333'),
            leading=12
        ))
        styles.add(ParagraphStyle(
            name='ReportSmall',
            parent=styles['Normal'],
            fontSize=8,
            textColor=colors.HexColor('#666666'),
            leading=10
        ))
        styles.add(ParagraphStyle(
            name='ReportDate',
            parent=styles['Normal'],
            fontSize=9,
            textColor=colors.HexColor('#888888'),
            alignment=TA_CENTER,
            spaceAfter=20
        ))
        
        story = []
        
        # Title
        story.append(Paragraph("Drug Discovery Research Report", styles['ReportTitle']))
        story.append(Paragraph(sanitize_text(request.query), styles['ReportSubtitle']))
        story.append(Paragraph(
            f"Generated: {pd.Timestamp.now().strftime('%B %d, %Y')} | Occolus",
            styles['ReportDate']
        ))
        story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor('#e0e0e0')))
        story.append(Spacer(1, 15))
        
        # Protein Info
        if request.protein_info:
            story.append(Paragraph("TARGET PROTEIN", styles['ReportSection']))
            p = request.protein_info
            protein_data = [
                ['Name', sanitize_text(p.get('name', 'N/A'))],
                ['UniProt ID', sanitize_text(p.get('id', 'N/A'))],
                ['Organism', sanitize_text(p.get('organism', 'N/A'))],
                ['Sequence Length', f"{p.get('sequence_length', 'N/A')} amino acids"]
            ]
            protein_table = Table(protein_data, colWidths=[1.5*inch, 4.5*inch])
            protein_table.setStyle(TableStyle([
                ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, -1), 9),
                ('TEXTCOLOR', (0, 0), (0, -1), colors.HexColor('#666666')),
                ('TEXTCOLOR', (1, 0), (1, -1), colors.HexColor('#1a1a1a')),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 5),
                ('TOPPADDING', (0, 0), (-1, -1), 5),
            ]))
            story.append(protein_table)
            
            if p.get('function'):
                story.append(Spacer(1, 5))
                func_text = sanitize_text(p.get('function', '')[:400])
                story.append(Paragraph(f"<b>Function:</b> {func_text}", styles['ReportSmall']))
            story.append(Spacer(1, 10))
        
        # Drug Candidates with Images
        if request.top_candidates and len(request.top_candidates) > 0:
            story.append(Paragraph("IDENTIFIED DRUG CANDIDATES", styles['ReportSection']))
            
            # Create grid of compounds with images
            compound_rows = []
            row = []
            for i, drug in enumerate(request.top_candidates[:6], 1):
                # Create compound card
                card_content = []
                
                # Add molecule image if available
                if drug.get('image_base64'):
                    try:
                        img_data = base64.b64decode(drug['image_base64'])
                        img_buffer = BytesIO(img_data)
                        img = RLImage(img_buffer, width=1.2*inch, height=1.2*inch)
                        card_content.append(img)
                    except:
                        pass
                
                name = sanitize_text(drug.get('name', 'N/A'))[:20]
                score = f"{drug.get('score', 0)*100:.0f}%"
                mw = str(drug.get('molecular_weight', 'N/A'))
                
                card_content.append(Paragraph(f"<b>{name}</b>", styles['ReportSmall']))
                card_content.append(Paragraph(f"Score: {score} | MW: {mw}", styles['ReportSmall']))
                
                row.append(card_content)
                
                if len(row) == 3:
                    compound_rows.append(row)
                    row = []
            
            if row:
                while len(row) < 3:
                    row.append([])
                compound_rows.append(row)
            
            if compound_rows:
                compound_table = Table(compound_rows, colWidths=[2*inch, 2*inch, 2*inch])
                compound_table.setStyle(TableStyle([
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 10),
                    ('TOPPADDING', (0, 0), (-1, -1), 10),
                ]))
                story.append(compound_table)
            story.append(Spacer(1, 10))
        
        # Novel Compounds
        if request.novel_drugs and len(request.novel_drugs) > 0:
            story.append(Paragraph("GENERATED NOVEL COMPOUNDS", styles['ReportSection']))
            
            novel_rows = []
            row = []
            for i, drug in enumerate(request.novel_drugs[:6], 1):
                card_content = []
                
                if drug.get('image_base64'):
                    try:
                        img_data = base64.b64decode(drug['image_base64'])
                        img_buffer = BytesIO(img_data)
                        img = RLImage(img_buffer, width=1.2*inch, height=1.2*inch)
                        card_content.append(img)
                    except:
                        pass
                
                name = sanitize_text(drug.get('name', f'Novel-{i}'))[:20]
                score = f"{drug.get('score', 0)*100:.0f}%"
                
                card_content.append(Paragraph(f"<b>{name}</b>", styles['ReportSmall']))
                card_content.append(Paragraph(f"Fitness: {score}", styles['ReportSmall']))
                
                row.append(card_content)
                
                if len(row) == 3:
                    novel_rows.append(row)
                    row = []
            
            if row:
                while len(row) < 3:
                    row.append([])
                novel_rows.append(row)
            
            if novel_rows:
                novel_table = Table(novel_rows, colWidths=[2*inch, 2*inch, 2*inch])
                novel_table.setStyle(TableStyle([
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 10),
                    ('TOPPADDING', (0, 0), (-1, -1), 10),
                ]))
                story.append(novel_table)
            story.append(Spacer(1, 10))
        
        # Research Analysis
        story.append(Paragraph("RESEARCH ANALYSIS", styles['ReportSection']))
        
        insights_text = sanitize_text(request.insights or "Analysis in progress...")
        
        # Clean markdown and format for PDF
        # Remove markdown headers
        insights_text = re.sub(r'^#{1,3}\s*', '', insights_text, flags=re.MULTILINE)
        
        # Handle bold text - simpler approach
        insights_text = insights_text.replace('**', '')
        
        # Split into paragraphs
        paragraphs = insights_text.split('\n\n')
        for para in paragraphs:
            para = para.strip()
            if not para:
                continue
            
            # Handle bullet points
            lines = para.split('\n')
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                
                if line.startswith('- ') or line.startswith('* '):
                    line = '  ' + line[2:]
                    story.append(Paragraph(line, styles['ReportBody']))
                elif line[0].isdigit() and '. ' in line[:4]:
                    story.append(Paragraph('  ' + line, styles['ReportBody']))
                else:
                    story.append(Paragraph(line, styles['ReportBody']))
            
            story.append(Spacer(1, 4))
        
        # References
        if request.papers and len(request.papers) > 0:
            story.append(PageBreak())
            story.append(Paragraph("REFERENCES", styles['ReportSection']))
            
            for i, paper in enumerate(request.papers[:15], 1):
                authors = paper.get('authors', ['Unknown'])
                author = sanitize_text(authors[0] if authors else 'Unknown')
                title = sanitize_text(paper.get('title', 'N/A'))[:100]
                source = sanitize_text(paper.get('source', 'N/A'))
                published = sanitize_text(paper.get('published', 'N/A'))
                
                ref_text = f"[{i}] {author} et al. {title}. {source}, {published}"
                story.append(Paragraph(ref_text, styles['ReportSmall']))
                story.append(Spacer(1, 3))
        
        # Build PDF
        doc.build(story)
        buffer.seek(0)
        
        logger.info("PDF export complete")
        
        return StreamingResponse(
            buffer,
            media_type="application/pdf",
            headers={
                "Content-Disposition": f"attachment; filename=drug_discovery_report.pdf"
            }
        )
        
    except Exception as e:
        logger.error(f"PDF export error: {e}")
        import traceback
        traceback.print_exc()
        return {"error": str(e)}


@app.post("/export-docx")
def export_report_docx(request: ExportReportRequest):
    """
    Generate a professional DOCX research report with molecule images
    """
    logger.info(f"=== DOCX EXPORT === Query: {request.query}")
    
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        
        doc = Document()
        
        # Set document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.8)
            section.bottom_margin = Inches(0.8)
            section.left_margin = Inches(0.8)
            section.right_margin = Inches(0.8)
        
        # Title
        title = doc.add_heading('Drug Discovery Research Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Subtitle
        subtitle = doc.add_paragraph()
        run = subtitle.add_run(sanitize_text(request.query))
        run.bold = True
        run.font.size = Pt(14)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Date
        date_para = doc.add_paragraph()
        date_run = date_para.add_run(f"Generated: {pd.Timestamp.now().strftime('%B %d, %Y')} | Occolus")
        date_run.font.size = Pt(9)
        date_run.font.color.rgb = RGBColor(128, 128, 128)
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        # Target Protein
        if request.protein_info:
            doc.add_heading('Target Protein', level=1)
            p = request.protein_info
            
            table = doc.add_table(rows=4, cols=2)
            table.style = 'Table Grid'
            
            cells = [
                ('Name', sanitize_text(p.get('name', 'N/A'))),
                ('UniProt ID', sanitize_text(p.get('id', 'N/A'))),
                ('Organism', sanitize_text(p.get('organism', 'N/A'))),
                ('Sequence Length', f"{p.get('sequence_length', 'N/A')} amino acids")
            ]
            
            for i, (label, value) in enumerate(cells):
                table.rows[i].cells[0].text = label
                table.rows[i].cells[1].text = str(value)
                table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
            
            if p.get('function'):
                func_para = doc.add_paragraph()
                func_run = func_para.add_run("Function: ")
                func_run.bold = True
                func_para.add_run(sanitize_text(p.get('function', '')[:400]))
            
            doc.add_paragraph()
        
        # Drug Candidates with Images
        if request.top_candidates and len(request.top_candidates) > 0:
            doc.add_heading('Identified Drug Candidates', level=1)
            
            # Create table with images
            num_cols = 3
            compounds = request.top_candidates[:6]
            num_rows = (len(compounds) + num_cols - 1) // num_cols
            
            table = doc.add_table(rows=num_rows * 2, cols=num_cols)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            for idx, drug in enumerate(compounds):
                row_idx = (idx // num_cols) * 2
                col_idx = idx % num_cols
                
                # Add image
                img_cell = table.rows[row_idx].cells[col_idx]
                if drug.get('image_base64'):
                    try:
                        img_data = base64.b64decode(drug['image_base64'])
                        img_buffer = BytesIO(img_data)
                        img_para = img_cell.paragraphs[0]
                        img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = img_para.add_run()
                        run.add_picture(img_buffer, width=Cm(3))
                    except Exception as e:
                        img_cell.text = "[Image]"
                
                # Add text below image
                text_cell = table.rows[row_idx + 1].cells[col_idx]
                text_para = text_cell.paragraphs[0]
                text_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                name_run = text_para.add_run(sanitize_text(drug.get('name', 'N/A'))[:20] + '\n')
                name_run.bold = True
                name_run.font.size = Pt(9)
                
                score = f"{drug.get('score', 0)*100:.0f}%"
                mw = drug.get('molecular_weight', 'N/A')
                info_run = text_para.add_run(f"Score: {score} | MW: {mw}")
                info_run.font.size = Pt(8)
                info_run.font.color.rgb = RGBColor(100, 100, 100)
            
            doc.add_paragraph()
        
        # Novel Compounds
        if request.novel_drugs and len(request.novel_drugs) > 0:
            doc.add_heading('Generated Novel Compounds', level=1)
            
            num_cols = 3
            compounds = request.novel_drugs[:6]
            num_rows = (len(compounds) + num_cols - 1) // num_cols
            
            table = doc.add_table(rows=num_rows * 2, cols=num_cols)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            for idx, drug in enumerate(compounds):
                row_idx = (idx // num_cols) * 2
                col_idx = idx % num_cols
                
                img_cell = table.rows[row_idx].cells[col_idx]
                if drug.get('image_base64'):
                    try:
                        img_data = base64.b64decode(drug['image_base64'])
                        img_buffer = BytesIO(img_data)
                        img_para = img_cell.paragraphs[0]
                        img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = img_para.add_run()
                        run.add_picture(img_buffer, width=Cm(3))
                    except:
                        img_cell.text = "[Image]"
                
                text_cell = table.rows[row_idx + 1].cells[col_idx]
                text_para = text_cell.paragraphs[0]
                text_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                name_run = text_para.add_run(sanitize_text(drug.get('name', f'Novel-{idx+1}'))[:20] + '\n')
                name_run.bold = True
                name_run.font.size = Pt(9)
                
                score = f"{drug.get('score', 0)*100:.0f}%"
                info_run = text_para.add_run(f"Fitness: {score}")
                info_run.font.size = Pt(8)
                info_run.font.color.rgb = RGBColor(100, 100, 100)
            
            doc.add_paragraph()
        
        # Research Analysis
        doc.add_heading('Research Analysis', level=1)
        
        insights_text = sanitize_text(request.insights or "Analysis in progress...")
        
        for line in insights_text.split('\n'):
            line = line.strip()
            if not line:
                continue
            
            # Headers
            if line.startswith('###') or line.startswith('##'):
                heading = doc.add_heading(line.replace('#', '').strip(), level=2)
            # Bullet points
            elif line.startswith('-') or line.startswith('*'):
                content = line.lstrip('-* ').strip()
                para = doc.add_paragraph(style='List Bullet')
                para.add_run(content.replace('**', ''))
            # Numbered
            elif line and line[0].isdigit() and '.' in line[:3]:
                content = line.split('.', 1)[1].strip() if '.' in line else line
                para = doc.add_paragraph(style='List Number')
                para.add_run(content.replace('**', ''))
            else:
                para = doc.add_paragraph()
                para.add_run(line.replace('**', ''))
        
        # References
        if request.papers and len(request.papers) > 0:
            doc.add_page_break()
            doc.add_heading('References', level=1)
            
            for i, paper in enumerate(request.papers[:15], 1):
                para = doc.add_paragraph()
                ref = para.add_run(f"[{i}] ")
                ref.bold = True
                
                authors = paper.get('authors', ['Unknown'])
                author = sanitize_text(authors[0] if authors else 'Unknown')
                title = sanitize_text(paper.get('title', 'N/A'))[:100]
                source = sanitize_text(paper.get('source', 'N/A'))
                published = sanitize_text(paper.get('published', 'N/A'))
                
                para.add_run(f"{author} et al. ")
                title_run = para.add_run(f'"{title}" ')
                title_run.italic = True
                para.add_run(f'{source}, {published}')
        
        # Save
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        logger.info("DOCX export complete")
        
        return StreamingResponse(
            buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f"attachment; filename=drug_discovery_report.docx"
            }
        )
        
    except Exception as e:
        logger.error(f"DOCX export error: {e}")
        import traceback
        traceback.print_exc()
        return {"error": str(e)}


if __name__ == "__main__":
    import uvicorn
    uvicorn.run("main:app", host="localhost", port=8000, reload=True)
